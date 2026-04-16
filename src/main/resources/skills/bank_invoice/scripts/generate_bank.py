#!/usr/bin/env python3
"""
XX银行 Excel/Word 生成脚本
基于模板生成新的 Excel 和 Word 文件，姓名自动脱敏
"""

import argparse
import sys
from datetime import datetime
from pathlib import Path

import openpyxl
from openpyxl import load_workbook
from docx import Document


def desensitize_name(name: str) -> str:
    """
    姓名脱敏处理
    - 两个字: X某 (如: 张三 -> 张某)
    - 三个字: XX某 (如: 张三丰 -> 张某某)
    - 四个及以上: XXX..某 (如: 张三丰四 -> 张某某某)
    """
    if not name:
        return "某某某"

    name = name.strip()
    if len(name) == 1:
        return f"{name}某"

    # 保留第一个字（姓氏），后面全部用"某"替换
    return name[0] + "某" * (len(name) - 1)


def generate_filename(name: str, serial: str, date_str: str = None, ext: str = "xlsx") -> str:
    """
    生成文件名: XX银行_脱敏姓名_YYMMDD_流水号.扩展名
    """
    desensitized = desensitize_name(name)

    if date_str is None:
        date_str = datetime.now().strftime("%y%m%d")

    return f"XX银行_{desensitized}_{date_str}_{serial}.{ext}"


def generate_excel(
    template_path: str,
    output_path: str,
    data: dict,
    date_str: str = None
) -> str:
    """
    基于模板生成新的 Excel 文件
    """
    # 加载模板
    wb = load_workbook(template_path)
    ws = wb.active

    # 数据列映射
    field_mapping = {
        'name': 'A',       # 姓名
        'id_card': 'B',    # 身份证号
        'phone': 'C',      # 联系电话
        'contract': 'D',   # 合同号
        'loan': 'E',       # 借据号
        'date': 'F',       # 放款日期
        'amount': 'G',     # 放款总金额
        'bank_amount': 'H',  # 行方承贷金额
        'fee_type': 'I',   # 费用类型
        'invoice': 'J',    # 发票金额
        'email': 'K',      # 邮箱
    }

    # 填充数据到第2行（第1行是表头）
    for field, col in field_mapping.items():
        if field in data and data[field] is not None:
            cell = ws[f"{col}2"]
            cell.value = data[field]

    # 生成文件名
    name = data.get('name', '某某某')
    serial = data.get('serial', datetime.now().strftime("%H%M%S"))
    filename = generate_filename(name, serial, date_str, "xlsx")

    # 确定输出路径
    output_dir = Path(output_path) if output_path else Path.cwd()
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / filename

    # 保存文件
    wb.save(output_file)
    return str(output_file)


def generate_word(
    template_path: str,
    output_path: str,
    data: dict,
    date_str: str = None
) -> str:
    """
    基于模板生成新的 Word 文件
    """
    # 加载模板
    doc = Document(template_path)

    # 获取数据
    name = data.get('name', '')
    phone = data.get('phone', '')
    email = data.get('email', '')
    id_card = data.get('id_card', '')
    invoice = data.get('invoice', 0)
    fee_type = data.get('fee_type', '')
    serial = data.get('serial', datetime.now().strftime("%H%M%S"))

    # 申请时间（提交日期）
    apply_date = datetime.now().strftime("%Y-%m-%d")

    # 填充表格数据
    table = doc.tables[0]

    # [1][1] 名称（姓名）
    table.cell(1, 1).text = name

    # [1][3] 纳税人识别号（身份证号）
    table.cell(1, 3).text = id_card

    # [4][1] 联系人姓名及电话
    table.cell(4, 1).text = f"{name} {phone}"

    # [4][3] 发票推送邮箱
    table.cell(4, 3).text = email

    # [8][1] 发票金额
    table.cell(8, 1).text = str(invoice)

    # [9][1] 发票摘要（使用费用类型）
    table.cell(9, 1).text = fee_type

    # [10][0] 发票附言（留白，不覆盖）

    # 填充提交日期（在第一个段落中）
    for para in doc.paragraphs:
        if "提交日期：" in para.text:
            # 找到提交日期所在的run并更新
            for run in para.runs:
                if "提交日期：" in run.text:
                    run.text = f"提交日期：{apply_date}"
                    break
            break

    # 生成文件名
    filename = generate_filename(name, serial, date_str, "docx")

    # 确定输出路径
    output_dir = Path(output_path) if output_path else Path.cwd()
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / filename

    # 保存文件
    doc.save(output_file)
    return str(output_file)


def parse_args():
    parser = argparse.ArgumentParser(
        description="基于模板生成XX银行 Excel 和 Word 文件（所有参数均为必填）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python generate_bank.py \\
    --name "张三丰" \\
    --id-card "123456200001011234" \\
    --phone "13800138000" \\
    --email "zhang@example.com" \\
    --contract "HT20240410001" \\
    --loan "JD20240410001" \\
    --date "2024-04-10" \\
    --amount 100000 \\
    --bank-amount 80000 \\
    --fee-type "服务费" \\
    --invoice 5000 \\
    --serial "001"
        """
    )

    parser.add_argument(
        '--template-dir',
        default='.',
        help='模板文件目录 (默认: 当前目录)'
    )

    parser.add_argument('--name', '-n', required=True, help='姓名（必填，会自动脱敏）')
    parser.add_argument('--id-card', '-i', required=True, help='身份证号（必填）')
    parser.add_argument('--phone', '-p', required=True, help='联系电话（必填）')
    parser.add_argument('--email', '-e', required=True, help='邮箱（必填）')
    parser.add_argument('--contract', '-c', required=True, help='合同号（必填）')
    parser.add_argument('--loan', '-l', required=True, help='借据号（必填）')
    parser.add_argument('--date', '-d', required=True, help='放款日期 (格式: YYYY-MM-DD 或 YYYYMMDD，必填)')
    parser.add_argument('--amount', type=float, required=True, help='放款总金额（必填）')
    parser.add_argument('--bank-amount', type=float, required=True, help='行方承贷金额（必填）')
    parser.add_argument('--fee-type', required=True, help='费用类型（必填）')
    parser.add_argument('--invoice', type=float, required=True, help='发票金额（必填）')
    parser.add_argument('--serial', '-s', required=True, help='流水号（必填）')
    parser.add_argument('--file-date', help='文件名中的日期 YYMMDD（默认为今天）')
    parser.add_argument('--output', '-o', help='输出目录（默认为当前目录）')

    return parser.parse_args()


def main():
    args = parse_args()

    # 模板路径
    template_dir = Path(args.template_dir)
    excel_template = template_dir / 'bank_template.xlsx'
    word_template = template_dir / 'bank_template.docx'

    # 检查模板文件
    if not excel_template.exists():
        print(f"错误: Excel模板文件不存在: {excel_template}")
        sys.exit(1)
    if not word_template.exists():
        print(f"错误: Word模板文件不存在: {word_template}")
        sys.exit(1)

    # 处理日期格式
    date_value = None
    if args.date:
        date_str = args.date.replace('-', '')
        try:
            date_obj = datetime.strptime(date_str, "%Y%m%d")
            date_value = date_obj.strftime("%Y-%m-%d")
        except ValueError:
            print(f"错误: 日期格式不正确: {args.date}")
            sys.exit(1)

    # 构建数据字典
    data = {
        'name': args.name,
        'id_card': args.id_card,
        'phone': args.phone,
        'email': args.email,
        'contract': args.contract,
        'loan': args.loan,
        'date': date_value,
        'amount': args.amount,
        'bank_amount': args.bank_amount,
        'fee_type': args.fee_type,
        'invoice': args.invoice,
        'serial': args.serial,
    }

    # 生成文件
    try:
        excel_file = generate_excel(
            template_path=str(excel_template),
            output_path=args.output,
            data=data,
            date_str=args.file_date
        )
        print(f"✓ Excel: {excel_file}")

        word_file = generate_word(
            template_path=str(word_template),
            output_path=args.output,
            data=data,
            date_str=args.file_date
        )
        print(f"✓ Word:  {word_file}")

        print(f"\n姓名脱敏: {args.name} -> {desensitize_name(args.name)}")
        print(f"申请时间: {datetime.now().strftime('%Y-%m-%d')}")

    except Exception as e:
        print(f"生成文件失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
